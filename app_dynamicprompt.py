import streamlit as st
import math
from io import BytesIO

from google.cloud import aiplatform

from PyPDF2 import PdfReader
from docx import Document
from fpdf import FPDF

PROJECT_ID = "7676662641783209984"
ENDPOINT_ID = "5504696168367521792"
LOCATION = "asia-south1"
try:
    aiplatform.init(project=PROJECT_ID, location=LOCATION)
except Exception as e:
    st.error(f"Could not initialize Vertex AI SDK. Have you run 'gcloud auth application-default login'? Error: {e}")
    st.stop()


st.set_page_config(
    page_title="Language Translation Agent",
    page_icon="🗺️",
    layout="wide",
    initial_sidebar_state="collapsed"
)

DEFAULT_PROMPT_INSTRUCTIONS = "You are a highly skilled translation expert. Your task is to translate the provided English text into the specified target language. Your output must ONLY be the translated text itself, with no introductory phrases."

def translate_text(final_prompt):
    """
    Sends the final prompt to your deployed Gemma model on a Vertex AI Endpoint.
    """
    if not final_prompt:
        return ""
    try:
        gemma_endpoint = aiplatform.Endpoint(
            endpoint_name=f"projects/{PROJECT_ID}/locations/{LOCATION}/endpoints/{ENDPOINT_ID}"
        )

        instances = [{"prompt": final_prompt}]

        response = gemma_endpoint.predict(instances=instances)
        translated_text = response.predictions[0]['content']
        return translated_text.strip()

    except Exception as e:
        st.error(f"An error occurred during translation with the Vertex AI endpoint: {e}")
        return None

def extract_text_from_pdf(file):
    try:
        pdf_reader = PdfReader(file)
        return [page.extract_text() for page in pdf_reader.pages if page.extract_text()]
    except Exception as e:
        st.error(f"Error reading PDF file: {e}")
        return []

def extract_text_from_docx(file):
    try:
        doc = Document(file)
        return [para.text for para in doc.paragraphs if para.text.strip()]
    except Exception as e:
        st.error(f"Error reading DOCX file: {e}")
        return []

def extract_text_from_txt(file):
    try:
        content = file.getvalue().decode("utf-8")
        lines = content.splitlines()
        return ["\n".join(lines[i:i + 40]) for i in range(0, len(lines), 40)]
    except Exception as e:
        st.error(f"Error reading TXT file: {e}")
        return []

def create_pdf_from_text(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
    pdf.set_font("DejaVu", size=12)
    pdf.multi_cell(0, 10, text)
    return bytes(pdf.output(dest='S'))

def create_docx_from_text(text):
    doc = Document()
    for paragraph in text.split('\n'):
        doc.add_paragraph(paragraph)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def init_session_state():
    if "text_translation_result" not in st.session_state:
        st.session_state.text_translation_result = ""
    if "doc_translation_result" not in st.session_state:
        st.session_state.doc_translation_result = ""
    if "target_language" not in st.session_state:
        st.session_state.target_language = "Abkhaz"
    if "doc_info" not in st.session_state:
        st.session_state.doc_info = {"name": None, "type": None}

init_session_state()

st.title("🗺️ Language Translation Agent (Gemma Model)")
st.markdown("Translate English text or documents using your deployed Gemma model on Vertex AI.")

languages = ["Abkhaz", "Acehnese", "Acholi", "Afar", "Afrikaans", "Albanian", "Alur", "Amharic", "Arabic", "Armenian", "Assamese", "Avar", "Awadhi", "Aymara","Azerbaijani", "Balinese", "Baluchi","Bambara", "BaoulÃ©", "Bashkir","Basque","Batak Karo","Batak Simalungun","Batak Toba","Belarusian","Bemba","Bengali","Betawi","Bhojpuri","Bikol","Bosnian","Breton","Bulgarian","Buryat", "Cantonese","Catalan","Cebuano","Chamorro","Chechen","Chichewa","Chinese (Simplified)","Chinese (Traditional)","Chuukese","Chuvash","Corsican","Crimean Tatar (Cyrillic)","Crimean Tatar (Latin)","Croatian","Czech","Danish","Dari","Dhivehi","Dinka","Dogri","Dombe","Dutch","Dyula","Dzongkha","English","Esperanto","Estonian","Ewe","Faroese","Fijian","Filipino","Finnish","Fon","French","French (Canada)","Frisian","Friulian","Fulani", "Ga","Galician","Georgian","German","Greek","Guarani","Gujrati","Haitian Crele","Hakha Chin","Hausa","Hawaiian","Hebrew","Hiligaynon","Hindi","Hmong","Hungarian","Hunsrik","Iban","Icelandic","Igbo","Ilocano","Indonesian","Inuktut (Latin)","Inuktut (Syllabics)","Irish","Italian","Jamaican Patois","Japanese","Javanese","Jingpo","Kalaallisut","Kannada","Kanuri","Kapampangan","Kazakh","Khasi","Khmer","Kiga","Kikongo","Kinyarwanda","Kituba","Kokborok","Komi","Konkani","Korean","Krio","Kurdish (Kurmanji)","Kurdish (Sorani)","Kyrgyz",
"Lao","Latgalian","Latin","Latvian","Ligurian","Limburgish","Lingala","Lithuanian","Lombard","Luganda","Luo","Luxembourgish", "Macedonian","Madurese","Maithili","Makassar","Malagasy","Malay","Malay (Jawi)","Malayalam","Maltese","Mam","Manx","Maroi","Marathi","Marshallese","Marwadi","Mauitian Creole","Meadow Mari","Meiteilon (Manipuri)","Minang","Mizo","Mongolian","Myanmar (Burmese)","Nahuatl (Eastern Huasteca)","Ndau","Ndebele (South)","Nepalbhasa (Newari)","Nepali","NKo","Norwegian","Nuer","Occitan","Odia (Oriya)","Oromo","Ossetian","Pangasinan","Papiamento","Pashto","Persian","Polish","Portuguese (Brazil)","Portuguese (Portugal)","Punjabi (Gurmuchhi)","Punjabi (Shahmukhi)","Quechua","QÃŠÂ¼eqchiÃŠÂ¼","Romani","Romanian","Rundi","Russian","Sami (North)","Samoan","Sango","Sanskrit","Santali (Latin)","Santali (Ol Chiki)","Scots Gaelic","Sepedi","Serbian","Sesotho","Seychellois Creole","Shan","Shona","Sicilian","Silesian","Sindhi","Sinhala","Slovak","Slovenian","Somali","Spanish","Sundanese", "Susu","Swahili","Swati","Swedish","Tahitian","Tazik","Tamazight","Tamazight (Tifinagh)","Tamil","Tatr","Telugu","Tetum","Thai","Tibetan","Tigrinya","Tiv","Tok Pisin","Tongan","Tshiluba","Tsonga","Tswana","Tulu","Tumbuka","Turkish","Turkmen","Tuvan","Twi","Udmurt","Ukrainian","Urdu","Uyghur","Uzbek","Venda","Venetian","Vietnamese","Waray","Welsh","Wolof","Xhosa","Yakut","Yiddish","Yoruba","Yucatec Maya","Zapotec","Zulu"]
languages.sort()
st.session_state.target_language = st.selectbox(
    "Select Target Language (you can also type to search)",
    languages,
    index=languages.index(st.session_state.target_language)
)

with st.expander("Advanced Options: Customize AI Instructions"):
    custom_instructions = st.text_area(
        "Enter custom instructions to override the default behavior:",
        placeholder="e.g., Translate the following text into a formal, business-appropriate tone.",
        height=150
    )

st.divider()

col1, col2 = st.columns([1, 1], gap="large")

with col1:
    st.subheader("Translate Text")
    input_text = st.text_area("Enter English text to translate:", height=200, key="text_input")
    if st.button("Translate Text", use_container_width=True, type="primary"):
        if input_text:
            instructions = custom_instructions if custom_instructions.strip() else DEFAULT_PROMPT_INSTRUCTIONS

            final_prompt = f"""
            {instructions}

            Target Language: {st.session_state.target_language}

            --- ENGLISH TEXT START ---
            {input_text}
            --- ENGLISH TEXT END ---
            """

            with st.spinner("Translating with Gemma..."):
                translated_output = translate_text(final_prompt)
                if translated_output is not None:
                    st.session_state.text_translation_result = translated_output
                    st.session_state.doc_translation_result = ""
                    st.session_state.doc_info = {"name": None, "type": None}
        else:
            st.warning("Please enter some text to translate.")
    st.divider()
    if st.session_state.text_translation_result:
        st.text_area("Text Translation Result", value=st.session_state.text_translation_result, height=250, disabled=True)

with col2:
    st.subheader("Translate Document")
    uploaded_file = st.file_uploader("Upload a document", type=["pdf", "docx", "txt"], key="file_uploader")
    if st.button("Translate Document", use_container_width=True, type="primary"):
        if uploaded_file is not None:
            instructions = custom_instructions if custom_instructions.strip() else DEFAULT_PROMPT_INSTRUCTIONS

            st.session_state.doc_info['name'] = uploaded_file.name
            file_ext = uploaded_file.name.split('.')[-1].lower()
            st.session_state.doc_info['type'] = file_ext

            source_chunks = []
            CHUNK_SIZE = 2
            if file_ext == "pdf": source_chunks = extract_text_from_pdf(uploaded_file)
            elif file_ext == "docx": source_chunks = extract_text_from_docx(uploaded_file)
            elif file_ext == "txt": source_chunks = extract_text_from_txt(uploaded_file)

            if source_chunks:
                translated_chunks = []
                total_chunks = math.ceil(len(source_chunks) / CHUNK_SIZE)
                progress_bar = st.progress(0, text=f"Translating chunk 1 of {total_chunks}...")

                for i in range(0, len(source_chunks), CHUNK_SIZE):
                    chunk_group = "\n".join(source_chunks[i:i + CHUNK_SIZE])
                    current_chunk_number = (i // CHUNK_SIZE) + 1
                    progress_bar.progress(current_chunk_number / total_chunks, text=f"Translating chunk {current_chunk_number} of {total_chunks}...")

                    final_prompt = f"""
                    {instructions}

                    Target Language: {st.session_state.target_language}

                    --- ENGLISH TEXT START ---
                    {chunk_group}
                    --- ENGLISH TEXT END ---
                    """

                    translated_output = translate_text(final_prompt)
                    if translated_output is not None:
                        translated_chunks.append(translated_output)
                    else:
                        st.error(f"Failed to translate chunk {current_chunk_number}. Skipping.")

                st.session_state.doc_translation_result = "\n\n".join(translated_chunks)
                st.session_state.text_translation_result = ""
                progress_bar.progress(1.0, text="Translation complete!")
            else:
                st.warning("Could not extract text from the document.")
        else:
            st.warning("Please upload a document to translate.")
    st.divider()
    if st.session_state.doc_translation_result:
        st.text_area("Document Translation Result", value=st.session_state.doc_translation_result, height=250, disabled=True)
        file_type = st.session_state.doc_info['type']
        file_name = f"translated_{st.session_state.doc_info['name']}"
        download_data = None

        if file_type == 'pdf':
            download_data = create_pdf_from_text(st.session_state.doc_translation_result)
        elif file_type == 'docx':
            download_data = create_docx_from_text(st.session_state.doc_translation_result)
        elif file_type == 'txt':
            download_data = st.session_state.doc_translation_result.encode('utf-8')

        if download_data:
            st.download_button(label=f"Download Translated .{file_type}", data=download_data, file_name=file_name, use_container_width=True)

